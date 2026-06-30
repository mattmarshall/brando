#!/usr/bin/env python3
"""brando office — a branded Word document template (also imports into Google Docs).

Generalized from the fastverk docx tool: a titled cover with the mark, in-brand
headings, body copy, and a color table with shaded swatch cells. The brand
supplies its palette, font, name, body copy, and swatches via a JSON config.

CLI: docx_gen.py <config.json> <out.docx> <icon.png>

config.json:
{
  "name": "fastverk",
  "subtitle": "Brand — document template",
  "font": "Space Grotesk",
  "ink": "#15161A", "accent": "#B5781A", "tertiary": "#4A565A",
  "body": "The mark ... one parametric source.",
  "swatches": [["15161A","bg midnight"], ["ECE7DA","fg cream"],
               ["E0A33E","accent amber"], ["4A565A","tertiary slate"]],
  "on_dark_hexes": ["15161A", "4A565A"]   # swatch fills that need cream text
}
"""
import json
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _rgb(hexstr):
    h = hexstr.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _run(p, text, size, color, font, bold=False):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold
    return r


def _para(doc, text, size, color, font, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    _run(p, text, size, color, font, bold)
    return p


def _shade(cell, hexcolor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(shd)


def build(cfg, out, icon):
    name = cfg["name"]
    subtitle = cfg.get("subtitle", "Brand — document template")
    font = cfg.get("font", "Helvetica")
    INK = _rgb(cfg["ink"])
    ACCENT = _rgb(cfg["accent"])
    SLATE = _rgb(cfg["tertiary"])
    cream_hex = cfg.get("cream", "#ECE7DA")
    CREAM = _rgb(cream_hex)
    body_copy = cfg.get("body", "Every asset derives from one parametric source.")
    swatches = cfg.get("swatches", [])
    on_dark = set(cfg.get("on_dark_hexes", []))

    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = font
    normal.size = Pt(11)
    normal.color.rgb = INK

    # cover
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.add_run().add_picture(icon, width=Inches(1.4))
    _para(doc, name, 30, INK, font, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    _para(doc, subtitle, 13, SLATE, font, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)

    _para(doc, "The mark", 17, ACCENT, font, bold=True, after=4)
    _para(doc, body_copy, 11, INK, font, after=14)

    if swatches:
        _para(doc, "Color", 17, ACCENT, font, bold=True, after=6)
        table = doc.add_table(rows=1, cols=len(swatches))
        table.style = "Table Grid"
        for cell, (hexc, label) in zip(table.rows[0].cells, swatches):
            _shade(cell, hexc)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(18)
            on = CREAM if hexc.lstrip("#").upper() in {h.upper() for h in on_dark} else INK
            _run(para, "#" + hexc, 10, on, font, bold=True)
            cap = cell.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(cap, label, 8, on, font)

    doc.save(out)
    print("wrote", out)


def main():
    cfg = json.load(open(sys.argv[1]))
    build(cfg, sys.argv[2], sys.argv[3])


if __name__ == "__main__":
    main()
